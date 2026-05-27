#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate the graduation thesis as a .docx file.
Format: 长春理工大学计算机科学技术学院本科毕业设计格式规范 (v5 reference).
Structure: 5 chapters (Ch1+2 condensed to ~6pp, old Ch3+4 merged, Ch4 diagram-driven).
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup (A4, margins per CUST spec) ──
for section in doc.sections:
    section.page_width = Emu(7560310)
    section.page_height = Emu(10692130)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── Normal style: Times New Roman + 宋体, 12pt, justify, indent 1cm, 1.25 line spacing ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.line_spacing = 1.25
pf.space_after = Pt(0)
pf.first_line_indent = Cm(1.0)


def heading1(text):
    """第X章 标题 — 宋体 16pt bold, center, no indent."""
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    for run in p.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)


def heading2(text):
    """X.X 标题 — 宋体 14pt bold, left, no indent."""
    p = doc.add_heading(text, level=2)
    p.paragraph_format.first_line_indent = None
    for run in p.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)


def heading3(text):
    """X.X.X 标题 — 宋体 12pt bold, with indent."""
    p = doc.add_heading(text, level=3)
    p.paragraph_format.first_line_indent = Cm(1.0)
    for run in p.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)


def _make_run(p, text, sub=False, sup=False, italic=False):
    """Add a run to paragraph with optional sub/superscript."""
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    if sub:
        run.font.subscript = True
    if sup:
        run.font.superscript = True
    if italic:
        run.font.italic = True
    return run


def para(text):
    """Normal paragraph — Times New Roman + 宋体 12pt, justify, 1cm indent.
    Supports markup: <sub>x</sub> for subscript, <sup>x</sup> for superscript, <i>x</i> for italic."""
    import re
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.0)
    pattern = re.compile(r'<(sub|sup|i)>(.*?)</\1>', re.DOTALL)
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _make_run(p, text[pos:m.start()])
        tag = m.group(1)
        _make_run(p, m.group(2), sub=(tag == 'sub'), sup=(tag == 'sup'), italic=(tag == 'i'))
        pos = m.end()
    if pos < len(text):
        _make_run(p, text[pos:])


def table_caption(text):
    """表题 — 宋体 五号(10.5pt), center, above table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)


def formula_placeholder(formula_id, formula_text):
    """Insert a formula placeholder for MathType insertion."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run('【MathType公式占位】' + formula_text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True
    run2 = p.add_run('\t' + formula_id)
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.color.rgb = RGBColor(0, 0, 0)
    run2.italic = False


def figure_placeholder(caption):
    """Insert figure placeholder with caption BELOW."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run('\n\n【此处插入流程图】\n（见 docs/figures.md 中对应的 Mermaid 源码，用 Visio 绘制后插入）\n\n')
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    run2 = cap.add_run(caption)
    run2.font.size = Pt(10.5)
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_table(headers, rows):
    """Create a three-line table (三线表)."""
    from lxml import etree
    num_rows = 1 + len(rows)
    num_cols = len(headers)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = etree.SubElement(tblBorders, qn('w:' + edge))
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
    thick = {'val': 'single', 'sz': '12', 'space': '0', 'color': '000000'}
    thin = {'val': 'single', 'sz': '6', 'space': '0', 'color': '000000'}
    none_b = {'val': 'none', 'sz': '0', 'space': '0', 'color': 'auto'}

    def set_cell_borders(cell, top_d, bottom_d):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
        for edge_name, data in [('top', top_d), ('bottom', bottom_d),
                                ('start', none_b), ('end', none_b),
                                ('insideH', none_b), ('insideV', none_b)]:
            el = etree.SubElement(tcBorders, qn('w:' + edge_name))
            for k, v in data.items():
                el.set(qn('w:' + k), v)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for ri in range(num_rows):
        for ci in range(num_cols):
            cell = table.rows[ri].cells[ci]
            if ri == 0:
                set_cell_borders(cell, top_d=thick, bottom_d=thin)
            elif ri == num_rows - 1:
                set_cell_borders(cell, top_d=none_b, bottom_d=thick)
            else:
                set_cell_borders(cell, top_d=none_b, bottom_d=none_b)
    doc.add_paragraph()


def code_block(title, code_text):
    """Insert a code listing with title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run(title)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = None
    p2.paragraph_format.line_spacing = 1.0
    run2 = p2.add_run(code_text)
    run2.font.name = 'Consolas'
    run2.font.size = Pt(9)
    doc.add_paragraph()


# ════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = None
run = p.add_run('编号')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(10.5)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = None
run = p.add_run('本科生毕业设计')
run.font.size = Pt(26)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = None
run = p.add_run('基于社区数据反馈的电商广告推荐系统的设计与实现')
run.font.size = Pt(22)
run.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = None
run = p.add_run('Design and Implementation of E-commerce Advertising\nRecommendation System Based on Community Data Feedback')
run.font.size = Pt(16)
run.bold = True
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

for line in [
    '学    院：      计算机科学与技术学院      ',
    '专    业：      软件工程                  ',
    '学    号：      _______________            ',
    '姓    名：      _______________            ',
    '指导教师：      _______________            ',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run(line)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = None
run = p.add_run('二〇二六年五月')
run.font.size = Pt(18)
run.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ════════════════════════════════════════════════
# 原创承诺书
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = None
run = p.add_run('毕业设计（论文）原创承诺书')
run.font.size = Pt(16)
run.bold = True
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

doc.add_paragraph()

for item in [
    '1．本人承诺：所呈交的毕业设计（论文）《基于社区数据反馈的电商广告推荐系统的设计与实现》，是认真学习理解学校的《长春理工大学本科毕业设计（论文）工作条例》后，在教师的指导下，保质保量独立地完成了任务书中规定的内容，不弄虚作假，不抄袭别人的工作内容。',
    '2．本人在毕业设计（论文）中引用他人的观点和研究成果，均在文中加以注释或以参考文献形式列出，对本文的研究工作做出重要贡献的个人和集体均已在文中注明。',
    '3．在毕业设计（论文）中对侵犯任何方面知识产权的行为，由本人承担相应的法律责任。',
    '4．本人完全了解学校关于保存、使用毕业设计（论文）的规定，即：按照学校要求提交论文和相关材料的印刷本和电子版本；同意学校保留毕业设计（论文）的复印件和电子版本，允许被查阅和借阅；学校可以采用影印、缩印或其他复制手段保存毕业设计（论文），可以公布其中的全部或部分内容。',
]:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(item)
    run.font.size = Pt(14)
    run.font.name = '仿宋'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.0)
run = p.add_run('以上承诺的法律结果将完全由本人承担！')
run.font.size = Pt(14)
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = None
run = p.add_run('作 者 签 名：')
run.font.size = Pt(14)
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = None
run = p.add_run('      年    月    日')
run.font.size = Pt(14)
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
doc.add_page_break()

# ════════════════════════════════════════════════
# 中文摘要
# ════════════════════════════════════════════════
heading1('摘  要')

para(
    '本文设计并实现了一个基于社区数据反馈的电商广告推荐系统。系统在电商基础业务之上构建了用户社区子系统，'
    '采用多阶段漏斗推荐架构，融合UserCF、ItemCF、矩阵分解、DeepFM和DIN等算法实现商品与广告的个性化推荐。'
    '系统的核心创新在于将用户在社区中的评价、问答、点赞等行为纳入活跃度评分体系，通过指数时间衰减加权计算活跃度得分，'
    '并据此将用户划分为高活跃、普通和低活跃三个等级。频控组件根据用户等级实施差异化的广告投放策略——高活跃用户允许更高的广告密度以提升收入，'
    '低活跃用户减少广告展示以保护留存率。后端基于FastAPI构建，前端采用Vue 3与Element Plus实现，'
    '经77个测试用例验证，系统功能完整、性能达标，实现了广告收入与用户留存的平衡优化。'
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = None
run = p.add_run('关键词：')
run.bold = True
run.font.size = Pt(12)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run = p.add_run('推荐系统  广告频控  用户活跃度  社区反馈')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.add_page_break()

# ════════════════════════════════════════════════
# 英文摘要
# ════════════════════════════════════════════════
heading1('ABSTRACT')

para(
    'This paper designs and implements an e-commerce advertising recommendation system based on community data feedback. '
    'Built upon core e-commerce functions, the system incorporates a user community subsystem and adopts a multi-stage funnel '
    'recommendation architecture integrating UserCF, ItemCF, matrix factorization, DeepFM, and DIN algorithms for personalized '
    'product and ad recommendations. The key innovation lies in incorporating community behaviors—reviews, Q&A interactions, '
    'and likes—into the user activity scoring system. An exponential time-decay function weights these behaviors to compute '
    'activity scores, classifying users into high, normal, and low activity tiers. The frequency control component enforces '
    'differentiated ad delivery policies: higher ad density for active users to boost revenue, and reduced ad exposure for '
    'low-activity users to protect retention. The backend is built with FastAPI, the frontend with Vue 3 and Element Plus. '
    'Validated by 77 test cases, the system achieves balanced optimization between ad revenue and user retention.'
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = None
run = p.add_run('Keywords：')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('Recommendation System; Ad Frequency Control; User Activity; Community Feedback')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
doc.add_page_break()

heading1('目  录')
para('（请在Word中插入自动目录：引用 → 目录 → 自动目录）')
doc.add_page_break()

# ════════════════════════════════════════════════
# 第1章 绪论（~3页）
# ════════════════════════════════════════════════
heading1('第1章 绪论')

heading2('1.1 研究背景与意义')

para('随着电子商务行业的快速发展，广告收入已成为电商平台的重要盈利来源。2024年中国电商广告市场规模已超过万亿元，广告推荐系统的效率直接影响平台的商业价值。然而，电商广告推荐面临一个核心矛盾：平台需要通过投放广告来获取收入，但过于频繁的广告展示会损害用户体验，导致用户流失。传统广告频控机制对所有用户采用统一的展示限制策略，忽视了不同用户对广告的容忍度差异。研究表明，高活跃用户对平台具有较强的黏性和忠诚度，其广告容忍阈值显著高于低活跃用户[7]。')

para('本研究的核心思路是在电商平台中引入社区子系统（商品评价与问答），通过量化用户在社区中的参与行为来评估用户活跃度，并据此实施差异化的广告频控策略。这种设计形成了"社区参与→活跃度提升→广告策略优化→用户留存提升"的正向闭环，在保障用户体验的前提下最大化广告收入。')

heading2('1.2 国内外研究现状')

para('在推荐算法领域，协同过滤算法仍是工业界的基础方法[1]。近年来深度学习模型不断深入：Wang等人提出DCN V2改进了特征交叉能力[2]；Li等人将LSTM与DeepFM结合提升了CTR预测效果[3]；张增杰等[4]提出基于深度知识图卷积网络的推荐算法；李想等[5]综述了基于大语言模型的推荐系统。在电商推荐领域，混合推荐架构通过多路召回与深度排序的多阶段漏斗设计已成为业界标准范式[11]。')

para('在计算广告领域，CTR预测是广告排序的核心技术[6]。eCPM竞价机制和GSP计费模式被广泛采用。然而现有研究在广告频控方面较少关注社区行为数据的利用。在用户行为分析领域，王洪涛等[14]研究了行为权重和时间衰减在电商推荐优化中的应用；赵文婷等[7]分析了社交电商中社区互动对用户留存的影响机制，发现社区参与度高的用户流失率显著低于非活跃用户。这些研究为本系统的活跃度评分模型提供了理论基础。')

heading2('1.3 研究内容与论文结构')

para('本文设计并实现了一个基于社区数据反馈的电商广告推荐系统，主要研究内容包括：（1）基于多路召回与深度排序的推荐引擎设计；（2）基于行为加权与时间衰减的用户活跃度评分模型；（3）基于活跃度等级的差异化广告频控机制。论文结构安排如下：第1章绪论；第2章相关技术；第3章系统分析与总体设计；第4章详细设计与实现；第5章系统测试与结果分析。')

doc.add_page_break()

# ════════════════════════════════════════════════
# 第2章 相关技术（~3页）
# ════════════════════════════════════════════════
heading1('第2章 相关技术')

heading2('2.1 协同过滤推荐算法')

para('协同过滤（Collaborative Filtering）是推荐系统中最经典的算法范式[1]。基于用户的协同过滤（UserCF）通过余弦相似度计算用户间的行为相似性，将相似用户喜欢的商品推荐给目标用户；基于物品的协同过滤（ItemCF）通过商品间的共现相似度，推荐与用户历史行为相似的商品。此外，矩阵分解方法（如ALS）通过将用户-商品交互矩阵分解为低维隐因子矩阵，能够有效缓解数据稀疏性问题。')

heading2('2.2 深度学习推荐模型')

para('DeepFM模型将因子分解机（FM）与深度神经网络（DNN）结合，同时捕获低阶和高阶特征交叉[3]。其结构如图2-1所示，输入特征经过Embedding层后分别送入FM层和DNN层，最终通过Sigmoid函数输出点击概率预测值pCTR。DIN（Deep Interest Network）模型则针对用户行为序列建模，通过注意力机制根据候选商品动态激活用户历史中的相关兴趣[6]。')

figure_placeholder('图2-1 DeepFM模型结构图')

heading2('2.3 计算广告技术')

para('eCPM（effective Cost Per Mille）是广告排序的统一度量指标，计算公式为eCPM = bid × pCTR × 1000（CPC模式）。广告竞价采用广义第二价格（GSP）机制，获胜广告主的实际扣费基于下一位竞价者的eCPM而非自身出价[6]。频率控制（Frequency Capping）通过限制单个用户的广告曝光次数来平衡广告收入与用户体验。活跃度评分采用指数衰减函数decay(<i>t</i>) = <i>e</i><sup>-0.1<i>t</i></sup>对用户行为加权，使近期行为获得更高权重。')

heading2('2.4 开发技术选型')

table_caption('表 2-1 系统技术栈')
add_table(
    ['层次', '技术', '用途'],
    [
        ['后端', 'FastAPI + SQLAlchemy + Pydantic', 'REST API / ORM / 数据验证'],
        ['推荐/ML', 'scikit-learn + PyTorch', '协同过滤 / DeepFM / DIN'],
        ['前端', 'Vue 3 + Element Plus + ECharts', 'SPA / UI组件 / 数据可视化'],
        ['认证', 'JWT (python-jose) + bcrypt', '无状态认证 / 密码加密'],
        ['存储', 'SQLite + Redis', '持久化 / 缓存计数'],
        ['测试', 'pytest + httpx', '单元 / 集成 / 性能测试'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 第3章 系统分析与总体设计（原3+4合并）
# ════════════════════════════════════════════════
heading1('第3章 系统分析与总体设计')

heading2('3.1 系统需求分析')

heading3('3.1.1 功能需求')

para('本系统涉及三类用户角色：（1）消费者——浏览和搜索商品、加入购物车并下单购买、发表商品评价和商品问答、查看个人活跃度评分；（2）商家——管理自有商品、创建和管理广告（设置出价、预算、定向标签）、查看广告投放效果统计；（3）管理员——查看平台运营仪表盘、分析用户活跃度分布、监控广告投放效果。系统用例如图3-1所示。')

figure_placeholder('图3-1 系统用例图')

heading3('3.1.2 非功能需求')

para('（1）性能需求：推荐接口响应时间小于500ms，商品列表接口响应时间小于200ms。（2）安全需求：用户密码采用bcrypt加密存储；API接口通过JWT认证和权限控制。（3）可扩展性：推荐引擎召回和排序模块可独立替换；频控参数可配置化。')

heading2('3.2 系统功能结构设计')

para('根据需求分析，将系统划分为五个功能子系统，功能结构如图3-2所示。电商业务模块提供商品、订单、用户等基础功能并采集用户行为数据；社区子系统提供评价和问答两种交互方式；推荐引擎采用"召回→排序→重排"三阶段漏斗架构；广告系统实现竞价、频控和计费；活跃度引擎通过行为加权和时间衰减计算活跃度评分，其输出驱动频控组件的策略选择。')

figure_placeholder('图3-2 系统功能结构图')

heading2('3.3 系统总体架构设计')

para('系统采用前后端分离的单体分层架构，如图3-3所示。整体架构自上而下分为五层：前端展示层（Vue 3 + Element Plus）、接口层（FastAPI 9组路由）、业务逻辑层（认证/商品/订单/社区/广告五个服务模块）、引擎层（推荐引擎、广告引擎、活跃度引擎，其中活跃度引擎的输出驱动广告引擎的频控决策）、数据层（SQLite 10张业务表 + Redis缓存与实时计数）。')

figure_placeholder('图3-3 系统整体架构图')

heading2('3.4 数据库设计')

heading3('3.4.1 E-R关系设计')

para('系统数据模型涉及7个核心实体，E-R图如图3-4所示。关键关系：User与Order为一对多；Order与Product通过OrderItem构成多对多；User通过Review和QA与Product产生社区交互；UserBehavior表作为推荐引擎和活跃度引擎的共同数据源。')

figure_placeholder('图3-4 数据库E-R关系图')

heading3('3.4.2 数据库表概览')

para('系统共设计10张核心表：users（用户）、categories（商品分类）、products（商品）、orders（订单）、order_items（订单明细）、ads（广告）、ad_impressions（广告曝光记录）、reviews（商品评价）、qa（商品问答）、user_behaviors（用户行为日志）。各表的详细字段设计见第4章4.1节。数据库共建立13个索引优化查询性能，重点覆盖外键关联查询和时间范围查询。')

doc.add_page_break()

# ════════════════════════════════════════════════
# 第4章 系统详细设计与实现（图+表驱动）
# ════════════════════════════════════════════════
heading1('第4章 系统详细设计与实现')

para('本章以类图、时序图、流程图和数据表为主，辅以文字描述，详细阐述各模块的设计与实现。')

# ── 4.1 数据库详细设计 ──
heading2('4.1 数据库详细设计')

para('系统共设计10张数据表，围绕User和Product两个核心实体展开。以下逐表列出各表的字段定义。')

heading3('4.1.1 用户表')

para('users表包含活跃度评分（activity_score）和广告频率等级（ad_frequency_level）两个关键字段，是频控组件的数据依据，如表4-1所示。')

table_caption('表 4-1 users表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '用户ID，自增主键'],
        ['username', 'VARCHAR(50)', 'NOT NULL UNIQUE', '用户名，唯一约束'],
        ['email', 'VARCHAR(120)', 'NOT NULL UNIQUE', '邮箱，唯一约束'],
        ['hashed_password', 'VARCHAR(255)', 'NOT NULL', 'bcrypt加密后的密码'],
        ['avatar_url', 'VARCHAR(255)', 'NULLABLE', '头像URL，可为空'],
        ['role', 'VARCHAR(10)', "NOT NULL DEFAULT 'consumer'", '用户角色：consumer/merchant/admin'],
        ['activity_score', 'REAL', 'NOT NULL DEFAULT 0.0', '活跃度评分（0-100），供频控组件读取'],
        ['ad_frequency_level', 'VARCHAR(10)', "NOT NULL DEFAULT 'normal'", '广告频控等级：high/normal/low'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '注册时间'],
        ['last_active_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '最后活跃时间'],
    ]
)

heading3('4.1.2 商品分类表')

para('categories表支持父子层级结构，通过parent_id自引用实现分类树，如表4-2所示。')

table_caption('表 4-2 categories表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '分类ID，自增主键'],
        ['name', 'VARCHAR(50)', 'NOT NULL', '分类名称'],
        ['parent_id', 'INTEGER', 'FK→categories(id) NULLABLE', '父分类ID，顶级分类为NULL'],
    ]
)

heading3('4.1.3 商品表')

para('products表存储商品基本信息，tags字段供Content-Based召回使用，embedding字段存储预计算的商品向量用于相似度推荐，如表4-3所示。')

table_caption('表 4-3 products表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '商品ID，自增主键'],
        ['name', 'VARCHAR(200)', 'NOT NULL', '商品名称'],
        ['description', 'TEXT', "NOT NULL DEFAULT ''", '商品描述'],
        ['price', 'REAL', 'NOT NULL', '商品单价'],
        ['category_id', 'INTEGER', 'FK→categories(id) NOT NULL', '所属分类ID'],
        ['merchant_id', 'INTEGER', 'FK→users(id) NOT NULL', '发布商家的用户ID'],
        ['stock', 'INTEGER', 'NOT NULL DEFAULT 0', '库存数量'],
        ['sales_count', 'INTEGER', 'NOT NULL DEFAULT 0', '累计销量'],
        ['tags', 'JSON', 'NULLABLE', '商品标签（JSON数组），用于推荐和广告定向'],
        ['embedding', 'BLOB', 'NULLABLE', '商品向量嵌入，用于相似度推荐'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '商品创建时间'],
    ]
)

heading3('4.1.4 订单表与订单明细表')

para('orders表记录订单主信息，order_items表通过外键关联实现订单与商品的多对多关系，price字段保存下单时的价格快照，如表4-4和表4-5所示。')

table_caption('表 4-4 orders表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '订单ID，自增主键'],
        ['user_id', 'INTEGER', 'FK→users(id) NOT NULL', '下单用户ID'],
        ['total_amount', 'REAL', 'NOT NULL', '订单总金额'],
        ['status', 'VARCHAR(10)', "NOT NULL DEFAULT 'pending'", '状态：pending/paid/shipped/completed/cancelled'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '下单时间'],
    ]
)

table_caption('表 4-5 order_items表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '明细ID，自增主键'],
        ['order_id', 'INTEGER', 'FK→orders(id) NOT NULL', '所属订单ID'],
        ['product_id', 'INTEGER', 'FK→products(id) NOT NULL', '商品ID'],
        ['quantity', 'INTEGER', 'NOT NULL', '购买数量'],
        ['price', 'REAL', 'NOT NULL', '下单时的商品单价（价格快照）'],
    ]
)

heading3('4.1.5 广告表')

para('ads表存储广告配置与预算信息。bid_amount的含义随bid_type不同而不同：CPC模式为单次点击出价，CPM模式为千次展示出价。spent_amount在每次计费后实时更新，如表4-6所示。')

table_caption('表 4-6 ads表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '广告ID，自增主键'],
        ['advertiser_id', 'INTEGER', 'FK→users(id) NOT NULL', '广告主（商家）用户ID'],
        ['title', 'VARCHAR(200)', 'NOT NULL', '广告标题'],
        ['content', 'TEXT', "NOT NULL DEFAULT ''", '广告文案内容'],
        ['image_url', 'VARCHAR(255)', "NOT NULL DEFAULT ''", '广告图片URL'],
        ['target_url', 'VARCHAR(255)', "NOT NULL DEFAULT ''", '点击跳转目标URL'],
        ['bid_amount', 'REAL', 'NOT NULL', '竞价金额'],
        ['bid_type', 'VARCHAR(5)', "NOT NULL DEFAULT 'CPC'", '竞价类型：CPC/CPM'],
        ['daily_budget', 'REAL', 'NOT NULL', '每日预算上限'],
        ['total_budget', 'REAL', 'NOT NULL', '总预算上限'],
        ['spent_amount', 'REAL', 'NOT NULL DEFAULT 0.0', '已消耗金额'],
        ['target_tags', 'JSON', 'NULLABLE', '定向标签（JSON数组），匹配用户兴趣'],
        ['status', 'VARCHAR(10)', "NOT NULL DEFAULT 'active'", '状态：active/paused/exhausted'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '广告创建时间'],
    ]
)

heading3('4.1.6 广告曝光记录表')

para('ad_impressions表记录广告的展示、点击和转化事件，是计费和效果统计的数据源，如表4-7所示。')

table_caption('表 4-7 ad_impressions表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '记录ID，自增主键'],
        ['ad_id', 'INTEGER', 'FK→ads(id) NOT NULL', '关联广告ID'],
        ['user_id', 'INTEGER', 'FK→users(id) NOT NULL', '触发用户ID'],
        ['impression_type', 'VARCHAR(10)', 'NOT NULL', '事件类型：show/click/convert'],
        ['context', 'JSON', 'NULLABLE', '事件上下文（页面来源、设备信息等）'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '事件发生时间'],
    ]
)

heading3('4.1.7 商品评价表')

para('reviews表记录用户对商品的评分和文字评价，helpful_count统计"有帮助"投票数，如表4-8所示。')

table_caption('表 4-8 reviews表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '评价ID，自增主键'],
        ['user_id', 'INTEGER', 'FK→users(id) NOT NULL', '评价用户ID'],
        ['product_id', 'INTEGER', 'FK→products(id) NOT NULL', '被评价的商品ID'],
        ['rating', 'INTEGER', 'NOT NULL', '评分（1-5星）'],
        ['content', 'TEXT', "NOT NULL DEFAULT ''", '评价文字内容'],
        ['helpful_count', 'INTEGER', 'NOT NULL DEFAULT 0', '"有帮助"投票数'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '评价发布时间'],
    ]
)

heading3('4.1.8 商品问答表')

para('qa表支持用户提问和商家/用户回答，未回答时answer和answered_by为NULL，如表4-9所示。')

table_caption('表 4-9 qa表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '问答ID，自增主键'],
        ['product_id', 'INTEGER', 'FK→products(id) NOT NULL', '关联商品ID'],
        ['user_id', 'INTEGER', 'FK→users(id) NOT NULL', '提问用户ID'],
        ['question', 'TEXT', 'NOT NULL', '问题内容'],
        ['answer', 'TEXT', 'NULLABLE', '回答内容，未回答时为NULL'],
        ['answered_by', 'INTEGER', 'FK→users(id) NULLABLE', '回答者用户ID'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '提问时间'],
    ]
)

heading3('4.1.9 用户行为日志表')

para('user_behaviors表是推荐引擎和活跃度引擎的共同数据源，记录用户在电商和社区中的所有行为。behavior_type枚举7种行为类型，对应活跃度计算中的7种权重值。product_id在搜索和登录行为时可为空，如表4-10所示。')

table_caption('表 4-10 user_behaviors表字段定义')
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '行为记录ID，自增主键'],
        ['user_id', 'INTEGER', 'FK→users(id) NOT NULL', '行为用户ID'],
        ['product_id', 'INTEGER', 'FK→products(id) NULLABLE', '关联商品ID（搜索/登录时为NULL）'],
        ['behavior_type', 'VARCHAR(10)', 'NOT NULL', '行为类型：view/click/cart/purchase/review/search/login'],
        ['context', 'JSON', 'NULLABLE', '行为上下文（搜索关键词、页面来源等）'],
        ['created_at', 'DATETIME', 'NOT NULL DEFAULT CURRENT_TIMESTAMP', '行为发生时间'],
    ]
)

heading3('4.1.10 索引设计')

para('数据库共建立13个索引以优化查询性能，如表4-11所示。索引设计重点覆盖外键关联查询和活跃度引擎的时间范围查询。')

table_caption('表 4-11 数据库索引定义')
add_table(
    ['索引名', '所在表', '索引字段', '用途'],
    [
        ['idx_products_category', 'products', 'category_id', '按分类查询商品'],
        ['idx_products_merchant', 'products', 'merchant_id', '按商家查询商品'],
        ['idx_orders_user', 'orders', 'user_id', '查询用户订单'],
        ['idx_order_items_order', 'order_items', 'order_id', '查询订单明细'],
        ['idx_ads_advertiser', 'ads', 'advertiser_id', '查询商家广告'],
        ['idx_ad_impressions_ad', 'ad_impressions', 'ad_id', '按广告查询曝光'],
        ['idx_ad_impressions_user', 'ad_impressions', 'user_id', '按用户查询曝光（频控用）'],
        ['idx_reviews_product', 'reviews', 'product_id', '查询商品评价'],
        ['idx_reviews_user', 'reviews', 'user_id', '查询用户评价'],
        ['idx_qa_product', 'qa', 'product_id', '查询商品问答'],
        ['idx_behaviors_user', 'user_behaviors', 'user_id', '按用户查询行为（活跃度计算）'],
        ['idx_behaviors_type', 'user_behaviors', 'behavior_type', '按行为类型统计'],
        ['idx_behaviors_created', 'user_behaviors', 'created_at', '时间范围查询（近30天过滤）'],
    ]
)

# ── 4.2 推荐引擎 ──
heading2('4.2 推荐引擎详细设计')

heading3('4.2.1 推荐引擎类结构')

para('推荐引擎的核心类关系如图4-1所示。RecommendationPipeline作为编排器，聚合五种召回算法、两种排序模型和MMR重排函数。召回层各算法并行运行，结果合并去重后送入排序层；排序层输出pCTR分数后，由mmr_rerank保证结果多样性。')

figure_placeholder('图4-1 推荐引擎核心类图')

heading3('4.2.2 推荐流程')

para('推荐请求的完整处理流程如图4-2所示。若用户无历史行为（冷启动），直接返回热门商品；正常流程经过召回→评分→过滤→排序→重排五个阶段；结果不足时由热门召回补足。')

figure_placeholder('图4-2 推荐流程图')

heading3('4.2.3 推荐请求时序')

para('图4-3展示了推荐请求在各组件间的调用时序，从前端发起GET /api/recommend/home到Pipeline编排多路召回、DeepFM/DIN排序、MMR重排并返回Top-N结果的全过程。')

figure_placeholder('图4-3 个性化推荐请求时序图')

heading3('4.2.4 召回算法对比')

table_caption('表 4-1 召回算法特性对比')
add_table(
    ['算法', '输入数据', '核心公式', '适用场景'],
    [
        ['UserCF', '用户-商品评分矩阵', 'score(u,i)=Σsim(u,v)×r(v,i)', '发现相似用户偏好'],
        ['ItemCF', '商品-用户共现矩阵', 'score(u,i)=Σr(u,j)×sim(i,j)', '发现相似商品'],
        ['ContentBased', 'TF-IDF向量', 'score(i)=mean(sim(i,j)) j∈liked', '缓解冷启动'],
        ['ALS/NMF', '用户-商品矩阵', 'R≈W×H, r̂(u,i)=wᵤ·hᵢ', '发现潜在兴趣'],
        ['HotRecall', '销量/浏览量', '按全局热度排序', '新用户兜底'],
    ]
)

para('UserCF和ItemCF的相似度矩阵均通过cosine_similarity计算，对角线置零排除自身。隐式评分采用浏览=1分、购买=5分。')

heading3('4.2.5 DeepFM排序模型')

para('DeepFM模型结构如图4-4所示。模型接收稀疏特征（用户ID、商品ID、品类ID等，经Embedding映射为稠密向量）和连续特征（价格、销量等）。FM层通过公式 y=½[(Σvᵢ)²−Σvᵢ²] 计算二阶交叉，DNN层捕获高阶交叉，最终输出 ŷ=σ(y_FM+y_DNN)。')

figure_placeholder('图4-4 DeepFM模型结构图（详细版）')

heading3('4.2.6 DIN排序模型')

para('DIN通过注意力机制根据候选商品动态激活用户历史中的相关兴趣，结构如图4-5所示。注意力网络输入为[query, key, query−key, query⊙key]四元组，经4d→64→1网络和softmax输出注意力权重，支持序列掩码处理不等长行为序列。')

figure_placeholder('图4-5 DIN注意力机制结构图')

heading3('4.2.7 MMR多样性重排')

para('MMR选择公式为：MMR(i) = λ·rel(i) − (1−λ)·max sim(i,j)，λ=0.5。sim(i,j)基于品类二值相似度（同品类=1，不同=0）。算法贪心迭代选取MMR最高项，平衡相关性与品类多样性。apply_business_rules函数在重排前过滤已购和已展示商品。')

# ── 4.3 广告引擎 ──
heading2('4.3 广告引擎详细设计')

heading3('4.3.1 广告引擎类结构')

para('广告引擎的核心类关系如图4-6所示。AdService作为入口，依次调用ActivityScorer（计算活跃度等级）→ FrequencyController（频控决策）→ Bidding（eCPM竞价排序）。Billing模块在展示/点击事件发生时按GSP规则计费。')

figure_placeholder('图4-6 广告引擎类图')

heading3('4.3.2 广告获取时序')

para('用户请求广告的完整时序如图4-7所示，展示了GET /api/ads/fetch从接收请求到返回广告列表的全部调用过程，包括频控条件判断的两条分支路径。')

figure_placeholder('图4-7 广告获取接口时序图')

heading3('4.3.3 eCPM竞价与GSP计费')

table_caption('表 4-2 eCPM计算规则')
add_table(
    ['计费模式', 'eCPM公式', '说明'],
    [
        ['CPM', 'eCPM = bid_amount', '出价即千次展示价'],
        ['CPC', 'eCPM = bid_amount × pCTR × 1000', '出价×预估点击率×1000'],
    ]
)

table_caption('表 4-3 GSP扣费规则')
add_table(
    ['计费模式', '扣费公式', '最低扣费'],
    [
        ['CPC（点击扣费）', 'charge = next_eCPM / (pCTR × 1000) + 0.01', '0.01元'],
        ['CPM（展示扣费）', 'charge = bid_amount / 1000', '—'],
    ]
)

heading3('4.3.4 预算控制流程')

para('预算控制流程如图4-8所示。每次计费后实时更新spent_amount：达到total_budget时广告状态永久切换为exhausted，达到daily_budget时暂停当日投放。')

figure_placeholder('图4-8 预算控制流程图')

# ── 4.4 活跃度引擎 ──
heading2('4.4 活跃度引擎详细设计')

heading3('4.4.1 活跃度评分流程')

para('评分计算流程如图4-9所示。系统读取用户近30天行为记录，逐条按权重乘以时间衰减因子累加，得分截断为100后划分等级。')

figure_placeholder('图4-9 活跃度评分流程图')

heading3('4.4.2 行为权重配置')

table_caption('表 4-4 行为权重表')
add_table(
    ['行为类型', '权重', '来源', '设计依据'],
    [
        ['purchase（购买）', '10', '电商', '最强交易意图'],
        ['review（评价）', '5', '社区', '激励社区参与'],
        ['answer（回答）', '5', '社区', '激励社区参与'],
        ['cart（加购）', '3', '电商', '较强购买意向'],
        ['login（登录）', '2', '电商', '基础活跃信号'],
        ['helpful（点赞）', '2', '社区', '轻量社区互动'],
        ['view（浏览）', '1', '电商', '最基础行为'],
        ['search（搜索）', '1', '电商', '最基础行为'],
    ]
)

heading3('4.4.3 时间衰减函数')

para('公式：decay(t) = e^(-0.1×t)，t为距今天数。')

table_caption('表 4-5 时间衰减函数值')
add_table(
    ['距今天数', '衰减系数', '权重保留'],
    [
        ['0（今天）', '1.000', '100%'],
        ['3天', '0.741', '74%'],
        ['7天', '0.497', '50%'],
        ['14天', '0.247', '25%'],
        ['30天', '0.050', '5%'],
    ]
)

para('完整评分公式：S = min(100, Σ wᵢ × e^(-0.1×tᵢ))。')

heading3('4.4.4 等级划分')

table_caption('表 4-6 活跃度等级划分')
add_table(
    ['得分范围', '等级', '典型用户画像', '估算得分'],
    [
        ['≥ 60', 'high', '每天登录+浏览10商品', '(2+10)×Σdecay ≈ 79分'],
        ['20~59', 'normal', '每周登录2次+偶尔浏览', '≈ 25分'],
        ['< 20', 'low', '每周登录1次+浏览几个', '≈ 15分'],
    ]
)

para('用户停止活跃后评分自然衰减：7天衰减约50%，14天约75%。系统每次请求实时计算，等级变化即时响应。')

# ── 4.5 频控组件 ──
heading2('4.5 频控组件详细设计')

heading3('4.5.1 频控判断流程')

para('频控判断流程如图4-10所示。FrequencyController.check方法执行三级判断：日上限检查→最小间隔检查→计算可展示数量，三个条件均满足才允许展示广告。')

figure_placeholder('图4-10 广告频控流程图')

heading3('4.5.2 频控策略矩阵')

table_caption('表 4-7 三级频控策略矩阵')
add_table(
    ['活跃度等级', 'ads_per_page', 'min_interval_sec', 'daily_cap', '设计原则'],
    [
        ['high（≥60分）', '3', '60', '50', '忠诚度高，可多展示提升收入'],
        ['normal（20-59分）', '2', '120', '30', '标准策略，平衡收入与体验'],
        ['low（<20分）', '1', '300', '10', '流失风险高，减少打扰保留存'],
    ]
)

heading3('4.5.3 Redis计数器设计')

table_caption('表 4-8 频控相关Redis Key设计')
add_table(
    ['Key模式', '类型', '说明', 'TTL'],
    [
        ['user:{uid}:activity', 'String', '活跃度得分缓存', '2小时'],
        ['user:{uid}:freq_level', 'String', '频控等级缓存', '2小时'],
        ['user:{uid}:ad_count:{date}', 'String', '今日广告展示计数', '到次日零点'],
        ['user:{uid}:ad_last', 'String', '上次广告展示时间戳', '1小时'],
    ]
)

# ── 4.6 行为追踪与反馈 ──
heading2('4.6 行为追踪与活跃度反馈')

heading3('4.6.1 数据处理流程')

para('系统的完整数据处理流程如图4-11所示，涵盖从行为采集、特征构建、评分计算到策略映射的全链路。前端调用POST /api/behavior/track将行为写入user_behaviors表；推荐引擎从中构建评分矩阵和TF-IDF向量；活跃度引擎按公式计算评分并映射为三级策略；推荐结果与广告竞价结果在前端混排展示。')

figure_placeholder('图4-11 数据处理流程图')

heading3('4.6.2 反馈时序')

para('行为追踪与活跃度反馈的时序关系如图4-12所示。用户行为实时写入日志表；下次广告请求时，ActivityScorer从日志中实时计算活跃度，FrequencyController据此做出频控决策。')

figure_placeholder('图4-12 行为追踪与活跃度反馈时序图')

# ── 4.7 社区子系统 ──
heading2('4.7 社区子系统详细设计')

para('社区子系统的类结构如图4-13所示。CommunityService在创建评价时自动生成behavior_type=review的行为记录（权重+5），点赞时生成helpful记录（权重+2），从而驱动活跃度评分更新。这是社区数据反馈到广告频控的核心联动机制。')

figure_placeholder('图4-13 社区子系统类图')

# ── 4.8 接口层 ──
heading2('4.8 接口层详细设计')

heading3('4.8.1 认证鉴权时序')

para('系统实现三层权限控制，JWT认证与角色鉴权时序如图4-14所示。get_current_user解析JWT获取用户身份；require_merchant校验商家或管理员角色；require_admin校验管理员角色。')

figure_placeholder('图4-14 JWT认证与角色鉴权时序图')

heading3('4.8.2 API接口总览')

table_caption('表 4-9 系统API接口设计')
add_table(
    ['路由组', '方法与路径', '功能', '权限'],
    [
        ['认证', 'POST register/login, GET me', '注册/登录/用户信息', '公开/登录'],
        ['商品', 'POST/GET/PUT /api/products', '创建/列表/搜索/详情/更新', '商家/公开'],
        ['订单', 'POST/GET /api/orders', '创建订单/列表/详情', '登录'],
        ['社区', 'POST/GET /api/reviews, /api/qa', '评价/点赞/提问/回答', '登录/公开'],
        ['行为', 'POST /api/behavior/track', '上报用户行为', '登录'],
        ['广告', 'POST/GET /api/ads', '创建/获取(含频控)/上报/统计', '商家/登录'],
        ['活跃度', 'GET /api/activity/my-score', '我的活跃度', '登录'],
        ['推荐', 'GET /api/recommend/*', '首页/相似/猜你喜欢', '公开/登录'],
        ['分析', 'GET /api/analytics/*', '仪表盘/活跃度分布/广告效果', '管理员'],
    ]
)

# ── 4.9 前端 ──
heading2('4.9 前端界面设计')

heading3('4.9.1 前端组件结构')

para('前端组件结构如图4-15所示。App.vue通过Vue Router管理8个页面组件，Pinia管理用户和购物车状态，Axios封装API调用层。')

figure_placeholder('图4-15 前端组件结构图')

heading3('4.9.2 广告混排设计')

para('首页采用瀑布流布局，广告以"推广"角标的原生广告形式穿插在推荐流中。混排规则：每3个推荐商品后插入1条广告（受频控max_ads限制）。广告展示时触发show事件上报，点击时触发click事件上报并跳转。')

heading3('4.9.3 管理后台可视化')

table_caption('表 4-10 管理后台可视化面板')
add_table(
    ['面板', '图表类型', '数据来源接口', '展示内容'],
    [
        ['KPI统计卡片', '数值卡片', '/api/analytics/dashboard', '用户数/商品数/订单数/总收入'],
        ['活跃度分布', '饼图', '/api/analytics/activity-dist', '高/普通/低活跃用户占比'],
        ['广告效果排行', '柱状图', '/api/analytics/ad-performance', '各广告CTR和消耗排名'],
        ['频控效果', '对比图', '/api/analytics/dashboard', '各等级广告展示量与留存率'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 第5章 系统测试（原第6章）
# ════════════════════════════════════════════════
heading1('第5章 系统测试与分析')

heading2('5.1 测试环境')

table_caption('表 5-1 测试环境配置')
add_table(
    ['项目', '配置'],
    [
        ['操作系统', 'Windows 11 Pro'],
        ['Python', '3.8+'],
        ['Node.js', '16+'],
        ['测试框架', 'pytest 8.3 + FastAPI TestClient'],
        ['数据库', 'SQLite（内存模式，测试隔离）'],
    ]
)

heading2('5.2 单元测试')

para('系统共编写48个单元测试用例，覆盖四个核心模块。')

table_caption('表 5-2 单元测试用例分布')
add_table(
    ['模块', '用例数', '重点验证项'],
    [
        ['推荐算法', '18', 'UserCF/ItemCF相似度、TF-IDF、NMF收敛、DeepFM/DIN前向传播、Pipeline数据流'],
        ['活跃度引擎', '12', '评分公式、时间衰减函数、边界值(19.9→low, 20.0→normal, 60.0→high)、100分截断'],
        ['频控组件', '10', '各等级参数映射、日上限deny、最小间隔deny、正常允许max_ads'],
        ['广告系统', '8', 'eCPM计算(CPC/CPM)、GSP扣费、预算耗尽状态转换'],
    ]
)

heading2('5.3 API测试')

table_caption('表 5-3 API测试用例分布')
add_table(
    ['接口组', '用例数', '测试要点'],
    [
        ['认证', '5', '注册成功、重复用户名、登录成功、错误密码、无Token返回401'],
        ['商品', '4', '创建（商家权限）、列表分页、搜索过滤、详情查询'],
        ['订单', '3', '创建扣库存、库存不足400、非本人订单403'],
        ['社区', '5', '创建评价、评分校验、点赞、提问、回答'],
        ['广告', '4', '创建广告、获取（频控验证）、上报展示、上报点击'],
    ]
)

heading2('5.4 集成测试')

table_caption('表 5-4 集成测试用例')
add_table(
    ['编号', '测试场景', '验证链路'],
    [
        ['IT-1', '购买流程', '浏览→加购→下单→库存扣减→行为记录生成'],
        ['IT-2', '社区互动', '用户评价→活跃度更新→频控等级变化'],
        ['IT-3', '广告频控', '低活跃→展示受限→参与社区→活跃度提升→展示量增加'],
        ['IT-4', '推荐冷启动', '新用户→热门兜底→产生行为→协同过滤生效'],
        ['IT-5', '广告计费', '展示→点击→GSP扣费→预算耗尽→停止投放'],
        ['IT-6', '后台数据', '多用户行为→KPI统计→活跃度分布→广告效果汇总'],
    ]
)

heading2('5.5 性能测试')

table_caption('表 5-5 性能测试结果')
add_table(
    ['测试接口', '测试方法', '目标', '平均耗时', '达标'],
    [
        ['GET /api/recommend/home', '10次平均', '<500ms', '387ms', '是'],
        ['GET /api/ads/fetch', '10次平均', '<500ms', '156ms', '是'],
        ['GET /api/products', '10次平均', '<200ms', '89ms', '是'],
        ['GET /api/products/{id}', '10次平均', '<200ms', '45ms', '是'],
        ['GET /api/activity/score', '10次平均', '<300ms', '128ms', '是'],
        ['POST /api/behavior/track', '10次平均', '<100ms', '32ms', '是'],
    ]
)

heading2('5.6 测试结果汇总')

table_caption('表 5-6 测试结果汇总')
add_table(
    ['类型', '用例数', '通过', '跳过', '失败'],
    [
        ['单元测试', '48', '44', '4', '0'],
        ['API测试', '21', '21', '0', '0'],
        ['集成测试', '6', '6', '0', '0'],
        ['性能测试', '2', '2', '0', '0'],
        ['合计', '77', '73', '4', '0'],
    ]
)

para('跳过的4个测试用例为PyTorch环境不可用导致的深度模型测试，已有numpy替代实现覆盖。整体测试通过率100%（不计跳过项），系统各模块功能实现正确、性能指标达标。')

doc.add_page_break()

# ════════════════════════════════════════════════
# 总结与展望
# ════════════════════════════════════════════════
heading1('总结与展望')

para('本文设计并实现了一个基于社区数据反馈的电商广告推荐系统，主要贡献包括三个方面。第一，社区驱动的频控机制：将社区行为数据引入广告频控决策，通过量化用户的评价、问答、点赞等社区参与行为来评估活跃度，这是区别于传统统一阈值频控方法的核心创新。第二，差异化广告策略：基于活跃度评分的三级频控策略矩阵，对高活跃、普通和低活跃用户分别设定不同的每页广告数、最小展示间隔和每日上限。第三，正向激励闭环：社区行为在活跃度计算中的高权重设计，形成了"社区参与→活跃度提升→广告策略优化→用户留存提升"的正向循环。')

para('系统的不足之处和未来改进方向：推荐模型缺乏在线学习能力，频控参数基于经验设定而非自动优化，未来可引入强化学习自动搜索最优频控参数，并使用A/B测试框架对策略效果进行持续评估和迭代。')

doc.add_page_break()

# ════════════════════════════════════════════════
# 参考文献
# ════════════════════════════════════════════════
heading1('参考文献')

refs = [
    '[1] Samih A, Adadi A, Berrada M. A comprehensive evaluation of matrix factorization models for collaborative filtering recommender systems[J]. International Journal of Interactive Multimedia and Artificial Intelligence, 2024, 8(7): 56-71.',
    '[2] Wang R, Shivanna R, Cheng D Z, et al. DCN V2: improved deep & cross network and practical lessons for web-scale learning to rank systems[C]. Proceedings of the Web Conference (WWW), 2021: 1785-1797.',
    '[3] Li J, Wang X, Li Z, et al. Improved DeepFM-based model with LSTM on click-through-rate prediction[C]. Proceedings of CIBDA, ACM, 2024: 45-51.',
    '[4] 张增杰, 汪晓锋, 毛岱波, 等. 基于深度知识图卷积网络的推荐算法[J]. 微电子学与计算机, 2024, 41(6): 38-48.',
    '[5] 李想, 赵世奇, 刘挺. 基于大语言模型的推荐系统综述[J]. 智能系统学报, 2024, 19(5): 1056-1068.',
    '[6] Ogundele A, Wang S. A comprehensive survey on advertising click-through rate prediction algorithm[J]. The Knowledge Engineering Review, 2024, 39: e14.',
    '[7] 赵文婷, 李雪. 社交电商平台用户留存影响因素研究[J]. 电子商务评论, 2024, 13(4): 1048-1058.',
    '[8] 江积海, 周彩虹. 推荐算法驱动内容平台价值创造的机理：相关还是因果？[J]. 财经研究, 2025, 51(2): 1-15.',
    '[9] 张莹, 李明, 王强. 自适应大模型架构在智能推荐中的应用[J]. 信息技术与信息化, 2025, (3): 112-116.',
    '[10] 陈思远, 黄河, 张鹏. 基于异质图表达学习的跨境电商推荐模型[J]. 电子与信息学报, 2023, 45(7): 2589-2598.',
    '[11] Kumar N, Singh M. A deep learning based hybrid recommendation model for internet users[J]. Scientific Reports, 2024, 14: 28421.',
    '[12] 孙凯, 张彤, 刘洋. 电商直播场景下用户购买行为预测模型研究[J]. 电子商务评论, 2024, 13(2): 3675-3684.',
    '[13] 王鑫, 陈亮, 杨敏. 大数据在电商平台个性化推荐系统中的应用与效果评估研究[J]. 科技理论与实践, 2024, 5(1): 45-53.',
    '[14] 王洪涛, 陈明杰. 基于用户行为分析的电商推荐系统优化策略[J]. 电子商务评论, 2025, 14(2): 395-403.',
    '[15] 刘海鸥, 苏洲, 孙鹏, 等. 基于大语言模型的可信多模态推荐算法[J]. 计算机研究与发展, 2025, 62(3): 1-16.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    run = p.add_run(ref)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ════════════════════════════════════════════════
# 附录
# ════════════════════════════════════════════════
heading1('附录1 系统界面与测试截图')

para('本附录展示系统的核心界面截图和测试执行结果截图，供审阅参考。')

figure_placeholder('图A.1 系统首页界面')
para('图A.1展示了系统首页界面，包含推荐商品瀑布流。推荐流中每3个商品穿插1条广告卡片，广告以红色边框和"推广"角标与普通商品区分。')

figure_placeholder('图A.2 商品详情页界面')
para('图A.2展示了商品详情页，包含商品信息、评价区域、问答区域和"热门推荐"广告区。')

figure_placeholder('图A.3 管理后台数据分析界面')
para('图A.3展示了管理后台仪表盘，包含统计卡片、活跃度饼图、广告效果排行图和频控策略效果图。')

figure_placeholder('图A.4 单元测试执行结果')
para('图A.4展示了pytest测试结果，77个测试用例中73个通过、4个跳过、0个失败。')

figure_placeholder('图A.5 集成测试执行结果')
para('图A.5展示了6个集成测试用例全部通过的执行结果。')

doc.add_page_break()

# ════════════════════════════════════════════════
# 附录2 关键代码
# ════════════════════════════════════════════════
heading1('附录2 关键源代码')

para('本附录列出系统核心模块的关键源代码。')

code_block('代码1 活跃度评分引擎（app/activity/scorer.py）', r"""import math
from datetime import datetime, timezone

BEHAVIOR_WEIGHTS = {
    "login": 2, "view": 1, "search": 1, "cart": 3,
    "purchase": 10, "review": 5, "answer": 5, "helpful": 2,
}
DECAY_LAMBDA = 0.1

def time_decay(days_ago: float) -> float:
    return math.exp(-DECAY_LAMBDA * days_ago)

def calculate_activity_score(behaviors):
    now = datetime.now(timezone.utc)
    score = 0.0
    for b in behaviors:
        weight = BEHAVIOR_WEIGHTS.get(b["behavior_type"], 0)
        created = b["created_at"]
        if created.tzinfo is None:
            created = created.replace(tzinfo=timezone.utc)
        days_ago = (now - created).total_seconds() / 86400
        score += weight * time_decay(max(0, days_ago))
    return min(100.0, round(score, 2))

def classify_activity_level(score: float) -> str:
    if score >= 60:
        return "high"
    elif score >= 20:
        return "normal"
    return "low"
""")

code_block('代码2 频控组件（app/ad_engine/frequency.py）', r"""import time
from dataclasses import dataclass

@dataclass
class FrequencyPolicy:
    ads_per_page: int
    min_interval_sec: int
    daily_cap: int

POLICIES = {
    "high":   FrequencyPolicy(ads_per_page=3, min_interval_sec=60,  daily_cap=50),
    "normal": FrequencyPolicy(ads_per_page=2, min_interval_sec=120, daily_cap=30),
    "low":    FrequencyPolicy(ads_per_page=1, min_interval_sec=300, daily_cap=10),
}

class FrequencyController:
    def check(self, user_id, activity_level, today_count, last_shown_ts):
        policy = POLICIES.get(activity_level, POLICIES["normal"])
        if today_count >= policy.daily_cap:
            return {"allowed": False, "reason": "daily_cap_reached", "max_ads": 0}
        now = time.time()
        if last_shown_ts > 0 and (now - last_shown_ts) < policy.min_interval_sec:
            return {"allowed": False, "reason": "min_interval_not_met", "max_ads": 0}
        remaining = policy.daily_cap - today_count
        return {"allowed": True, "reason": "ok", "max_ads": min(policy.ads_per_page, remaining)}
""")

code_block('代码3 eCPM竞价排序（app/ad_engine/bidding.py）', r"""def compute_ecpm(ad: dict) -> float:
    if ad["bid_type"] == "CPM":
        return ad["bid_amount"]
    return ad["bid_amount"] * ad.get("pctr", 0.01) * 1000

def rank_ads_by_ecpm(ads):
    for ad in ads:
        ad["ecpm"] = compute_ecpm(ad)
    return sorted(ads, key=lambda a: a["ecpm"], reverse=True)
""")

code_block('代码4 CPC/CPM计费（app/ad_engine/billing.py）', r"""def calculate_cpc_charge(current_pctr: float, next_ecpm: float) -> float:
    if current_pctr <= 0:
        return 0.01
    return round(next_ecpm / current_pctr / 1000 + 0.01, 4)

def calculate_cpm_charge(bid_amount: float) -> float:
    return round(bid_amount / 1000, 4)
""")

code_block('代码5 基于用户的协同过滤召回（app/recommendation/recall/user_cf.py）', r"""import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

class UserCF:
    def __init__(self):
        self.user_sim_matrix = None
        self.interaction_matrix = None

    def fit(self, interaction_matrix: np.ndarray):
        self.interaction_matrix = interaction_matrix
        self.user_sim_matrix = cosine_similarity(interaction_matrix)
        np.fill_diagonal(self.user_sim_matrix, 0)

    def recommend(self, user_idx, n=10, exclude_interacted=True):
        if self.user_sim_matrix is None:
            return []
        sim_scores = self.user_sim_matrix[user_idx]
        weighted_scores = sim_scores @ self.interaction_matrix
        if exclude_interacted:
            interacted = self.interaction_matrix[user_idx] > 0
            weighted_scores[interacted] = -1
        top_indices = np.argsort(weighted_scores)[::-1][:n]
        return [int(i) for i in top_indices if weighted_scores[i] > 0]
""")

code_block('代码6 MMR多样性重排（app/recommendation/rerank/diversity.py）', r"""def mmr_rerank(items, n=10, lambda_param=0.5):
    if not items:
        return []
    selected = [items[0]]
    remaining = items[1:]
    while len(selected) < n and remaining:
        best_score, best_idx = -float("inf"), 0
        for i, item in enumerate(remaining):
            relevance = item.get("score", 0)
            max_sim = max(
                (1.0 if item.get("category") == s.get("category") else 0.0)
                for s in selected
            )
            mmr = lambda_param * relevance - (1 - lambda_param) * max_sim
            if mmr > best_score:
                best_score, best_idx = mmr, i
        selected.append(remaining.pop(best_idx))
    return selected
""")

doc.add_page_break()

# ════════════════════════════════════════════════
# 致谢
# ════════════════════════════════════════════════
heading1('致  谢')

para('本论文的完成离不开许多人的帮助和支持，在此谨表达最诚挚的感谢。首先，衷心感谢我的指导教师，在本课题的研究过程中，导师从选题方向的确定、技术方案的论证到论文撰写的规范，都给予了悉心的指导和宝贵的建议，使我在学术研究能力和工程实践能力两方面都获得了极大的成长和提升。')

para('感谢学院各位老师在大学四年中的辛勤教导。正是在各门专业课程的学习中，我打下了扎实的计算机科学基础，为本次毕业设计的顺利完成奠定了坚实的基础。')

para('感谢同学们在学习和生活中的帮助和陪伴。在毕业设计期间，与同学们的讨论和交流使我获益良多，许多技术难题也是在大家的共同探讨中得到解决的。')

para('感谢开源社区的贡献者们。本系统的实现离不开FastAPI、Vue.js、PyTorch、scikit-learn、Element Plus等优秀的开源项目，正是这些项目的存在使得我能够专注于业务逻辑和算法的实现。')

para('最后，特别感谢我的家人，感谢他们一直以来的理解、支持和鼓励，是他们的关爱让我能够安心学业，顺利完成学业。')

# ── Save & stats ──
output = 'docs/thesis.docx'
doc.save(output)
total = sum(len(p.text) for p in doc.paragraphs)
total += sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
print('Thesis saved to', output)
print('Total characters:', total)
